"""Generate the Ropener Curtain Controller user guide as a .docx file.

Source: Ropener-User-Guide.md (parsed at build time — single source of truth).
Run:  python make_user_guide.py
Output: "Ropener User Guide.docx" in the gitignored dist/ folder.
"""

import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Valar Systems branding (from the logo): black "VALAR" wordmark with an orange
# "systems" script. Headings stay black for readability; orange (#F89048) is
# the accent — divider rule, table header rows, and footer.
ACCENT = RGBColor(0x00, 0x00, 0x00)   # brand black (headings)
ORANGE = RGBColor(0xF8, 0x90, 0x48)   # Valar accent orange (sampled from logo)
ORANGE_HEX = "F89048"
GREY = RGBColor(0x59, 0x59, 0x59)     # muted grey for sub-text
MONO = "Consolas"
HERE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(HERE, "valar-logo-black.png")  # transparent bg, sits on white
DIST = os.path.join(HERE, "dist")     # generated output (gitignored)

doc = Document()


def shade_cell(cell, hex_fill):
    """Fill a table cell with a solid background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def hrule(color_hex=ORANGE_HEX, size=18):
    """Add a thin colored horizontal divider rule."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

# ---- base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    st = doc.styles[name]
    st.font.color.rgb = ACCENT
    st.font.size = Pt(size)


def code(text):
    """Add a paragraph of monospaced text (for POSIX strings / commands)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(10)
    return p


def callout(title, body, fill="FCE3CF"):
    """A shaded attention box (light orange) with a bold title and body text."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade_cell(cell, fill)
    cell.text = ""
    tp = cell.paragraphs[0]
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0x9C, 0x4A, 0x00)   # dark orange for contrast
    bp = cell.add_paragraph()
    bp.add_run(body)
    doc.add_paragraph()


def mono_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        shade_cell(hdr[i], ORANGE_HEX)          # brand-orange header row
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            # monospace any cell that looks like a POSIX string / value
            if any(ch in val for ch in (",", "<")) or val[:3] in (
                "EST", "CST", "MST", "PST", "AKS", "HST", "SST", "ChS", "AST", "UTC",
                "WET", "GMT", "CET", "EET", "MSK",
            ):
                run.font.name = MONO
                run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# ============================================================================
# TITLE / BRANDING
# ============================================================================
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO):
    logo_p.add_run().add_picture(LOGO, width=Inches(3.2))

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Ropener Curtain Controller")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("User Guide")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = ORANGE

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run("Firmware 2.6.1")
r.italic = True
r.font.color.rgb = GREY

hrule()              # orange divider under the title block
doc.add_paragraph()

# Branded footer on every page: company name + accent bullet + site.
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.text = ""
run = fp.add_run("Valar Systems")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = ACCENT
bullet = fp.add_run("   •   ")
bullet.bold = True
bullet.font.size = Pt(9)
bullet.font.color.rgb = ORANGE
url = fp.add_run("valarsystems.com")
url.font.size = Pt(9)
url.font.color.rgb = GREY


# ============================================================================
# CONTENT — parsed from Ropener-User-Guide.md (the single source of truth)
# ============================================================================
# The markdown is rendered with the branded helpers above. Document chrome (the
# title block and footer) and the table column widths stay here in the renderer,
# since they are presentation, not content.

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*|\[[^\]]+\]\([^)]+\))")

# Hand-tuned column widths (inches), keyed by the table's header row.
WIDTHS = {
    ("Button", "Action", "What it does"): [1.4, 1.8, 3.2],
    ("Control", "What to set"): [1.7, 4.7],
    ("Zone", "Where it's used", "With DST", "Without DST (standard all year)"): [1.3, 2.5, 1.6, 1.3],
    ("Setting", "What it controls", "Notes"): [1.4, 2.6, 2.4],
    ("Symptom", "What to try"): [2.1, 4.3],
    ("Control (web page)", "Purpose"): [2.4, 4.0],
    ("Physical button", "Action"): [2.4, 4.0],
}


def strip_inline(text):
    """Reduce inline markdown to plain text (drop emphasis, links -> their text)."""
    text = LINK_RE.sub(r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def render_inline(p, text):
    """Add runs to paragraph p, honoring **bold**, *italic*, `code`, [text](url)."""
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            p.add_run(tok[1:-1]).font.name = MONO
        elif tok.startswith("*"):
            p.add_run(tok[1:-1]).italic = True
        else:
            p.add_run(LINK_RE.match(tok).group(1))
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def render_table(block):
    def cells(row):
        return [strip_inline(c) for c in row.strip().strip("|").split("|")]
    header = cells(block[0])
    body = [cells(r) for r in block[2:]]          # block[1] is the --- separator
    mono_table(header, body, widths=WIDTHS.get(tuple(header)))


def render_quote(quoted):
    nonempty = [q for q in quoted if q]
    if nonempty and "⚠️" in nonempty[0]:                 # warning emoji -> callout box
        title = strip_inline(nonempty[0].replace("⚠️", ""))
        callout(title, strip_inline(" ".join(nonempty[1:])))
    else:                                          # short aside -> inline note
        render_inline(doc.add_paragraph(), " ".join(nonempty))


def render_markdown(path):
    with open(path, encoding="utf-8") as f:
        lines = f.read().split("\n")
    i, n, started, skip = 0, len(lines), False, False
    while i < n:
        line = lines[i].strip()
        if not line or line == "---":
            i += 1; continue
        if line.startswith("# "):                  # doc title -> chrome, skip
            i += 1; continue
        if line.startswith("## ") or line.startswith("### "):
            level = 2 if line.startswith("### ") else 1
            title = strip_inline(line[level + 1:])
            if title.lower() == "contents":        # the TOC is omitted in print
                skip = True; i += 1; continue
            skip = False
            doc.add_heading(title, level=level)
            started = True; i += 1; continue
        if skip:
            i += 1; continue
        if line.startswith("```"):
            i += 1; buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            code("\n".join(buf)); started = True; continue
        if line.startswith("|"):
            buf = []
            while i < n and lines[i].strip().startswith("|"):
                buf.append(lines[i]); i += 1
            render_table(buf); started = True; continue
        if line.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].strip()); i += 1
            render_quote(buf); started = True; continue
        if line.startswith("- "):
            while i < n and lines[i].strip().startswith("- "):
                render_inline(doc.add_paragraph(style="List Bullet"), lines[i].strip()[2:])
                i += 1
            started = True; continue
        if re.match(r"\d+\.\s", line):
            while i < n and re.match(r"\d+\.\s", lines[i].strip()):
                render_inline(doc.add_paragraph(style="List Number"),
                              re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            started = True; continue
        if line.startswith("*") and not line.startswith("**") and line.endswith("*"):
            if not started or line.lower().startswith("*valar"):
                i += 1; continue                   # firmware version / footer -> chrome
        render_inline(doc.add_paragraph(), line); started = True; i += 1


render_markdown(os.path.join(HERE, "Ropener-User-Guide.md"))

os.makedirs(DIST, exist_ok=True)
out = os.path.join(DIST, "Ropener User Guide.docx")
doc.save(out)
print("Saved:", out)
