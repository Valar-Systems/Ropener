"""Generate the Ropener Bill of Materials (V2.0) as a branded .docx file.

Source data: BOM-V2.0.md (parsed at build time — the single source of truth).
Run:  python make_bom.py
Output: "Ropener Bill of Materials.docx" in the gitignored dist/ folder.
"""

import os
import re

import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Valar Systems branding (matches the user guide).
ACCENT = RGBColor(0x00, 0x00, 0x00)   # brand black (headings)
ORANGE = RGBColor(0xF8, 0x90, 0x48)   # Valar accent orange
ORANGE_HEX = "F89048"
GREY = RGBColor(0x59, 0x59, 0x59)
LINK_BLUE = "0563C1"
MONO = "Consolas"
HERE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(HERE, "valar-logo-black.png")
DIST = os.path.join(HERE, "dist")     # generated output (gitignored)

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def parse_bom(md_path):
    """Parse BOM-V2.0.md into an ordered list of blocks:
      ("h1"|"h2", heading_text)  — from ## / ### markdown headings
      ("table", [(part, qty, desc), ...])  — from each markdown table

    The top-level "# " title and the "> " intro blockquote are document chrome
    handled by the renderer, so they are skipped here. Each table's first two
    rows (header + alignment separator) are dropped; the rest become data rows.
    Markdown links in the description cell are rendered as hyperlinks downstream.
    """
    blocks = []
    table_lines = []

    def flush_table():
        if not table_lines:
            return
        rows = []
        for ln in table_lines[2:]:                 # skip header + separator
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            if len(cells) >= 3:
                rows.append((cells[0], cells[1], cells[2]))
        if rows:
            blocks.append(("table", rows))
        table_lines.clear()

    with open(md_path, encoding="utf-8") as f:
        for raw in f:
            stripped = raw.strip()
            if stripped.startswith("|"):
                table_lines.append(stripped)
                continue
            flush_table()
            if stripped.startswith("### "):
                blocks.append(("h2", stripped[4:].strip()))
            elif stripped.startswith("## "):
                blocks.append(("h1", stripped[3:].strip()))
            # everything else (title, blockquote, blank lines) is skipped
        flush_table()
    return blocks


doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for name, size in (("Heading 1", 16), ("Heading 2", 12.5)):
    st = doc.styles[name]
    st.font.color.rgb = ACCENT
    st.font.size = Pt(size)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def hrule(color_hex=ORANGE_HEX, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_hyperlink(paragraph, url, text):
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def fill_desc(cell, text):
    """Write a description cell, turning [text](url) markdown into hyperlinks."""
    cell.text = ""
    p = cell.paragraphs[0]
    last = 0
    matched = False
    for m in LINK_RE.finditer(text):
        matched = True
        if m.start() > last:
            p.add_run(text[last:m.start()])
        add_hyperlink(p, m.group(2), m.group(1))
        last = m.end()
    if last < len(text):
        p.add_run(text[last:])
    if not matched and not text:
        p.add_run("")


def bom_table(rows, widths=(2.7, 0.7, 3.0)):
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Part", "Qty", "Description"]
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        shade_cell(c, ORANGE_HEX)
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        if i == 1:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for part, qty, desc in rows:
        cells = t.add_row().cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(part).bold = True
        cells[1].text = str(qty)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        fill_desc(cells[2], desc)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ---------------------------------------------------------------------------
# Header / branding
# ---------------------------------------------------------------------------
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO):
    logo_p.add_run().add_picture(LOGO, width=Inches(3.2))

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Ropener Curtain Controller")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Bill of Materials")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = ORANGE

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run("Version 2.0")
r.italic = True
r.font.color.rgb = GREY

hrule()
doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "Everything needed to build one Ropener. Quantities are per unit. Links in "
    "the Description column point to the custom PCB and suggested suppliers; "
    "equivalent parts may be substituted where noted."
)

# ---------------------------------------------------------------------------
# Body: sections + tables, parsed from BOM-V2.0.md (single source of truth)
# ---------------------------------------------------------------------------
for kind, payload in parse_bom(os.path.join(HERE, "BOM-V2.0.md")):
    if kind == "h1":
        doc.add_heading(payload, level=1)
    elif kind == "h2":
        doc.add_heading(payload, level=2)
    elif kind == "table":
        bom_table(payload)

# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.text = ""
run = fp.add_run("Valar Systems")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = ACCENT
bullet = fp.add_run("   •   ")
bullet.bold = True
bullet.font.size = Pt(9)
bullet.font.color.rgb = ORANGE
url = fp.add_run("valarsystems.com")
url.font.size = Pt(9)
url.font.color.rgb = GREY

os.makedirs(DIST, exist_ok=True)
out = os.path.join(DIST, "Ropener Bill of Materials.docx")
doc.save(out)
print("Saved:", out)
